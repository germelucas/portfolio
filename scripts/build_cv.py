from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "public" / "cv-lucas-germe.docx"

NAVY = RGBColor(16, 38, 66)
BLUE = RGBColor(35, 112, 181)
INK = RGBColor(30, 41, 59)
MUTED = RGBColor(82, 96, 115)


def set_run_font(run, size=9.4, color=INK, bold=False, italic=False):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_margins(cell, top=80, start=110, bottom=80, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        borders.append(element)


def configure_table_geometry(table, widths_dxa):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def add_hyperlink(paragraph, label, url, size=8.9):
    relationship = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    props.append(fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "526073")
    props.append(color)
    font_size = OxmlElement("w:sz")
    font_size.set(qn("w:val"), str(int(size * 2)))
    props.append(font_size)
    run.append(props)
    text = OxmlElement("w:t")
    text.text = label
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_separator(paragraph):
    run = paragraph.add_run("  |  ")
    set_run_font(run, size=8.8, color=MUTED)


def add_section_heading(doc, text):
    paragraph = doc.add_paragraph(style="CV Section")
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text.upper())
    set_run_font(run, size=10.5, color=BLUE, bold=True)
    return paragraph


def add_role(doc, title, organization, period, location=None):
    paragraph = doc.add_paragraph(style="CV Role")
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(title)
    set_run_font(run, size=10.0, color=NAVY, bold=True)
    run = paragraph.add_run(f" - {organization}")
    set_run_font(run, size=10.0, color=INK, bold=True)
    run = paragraph.add_run(f"    {period}")
    set_run_font(run, size=8.8, color=MUTED, bold=True)
    if location:
        run = paragraph.add_run(f" | {location}")
        set_run_font(run, size=8.8, color=MUTED)


def create_numbering(document):
    numbering = document.part.numbering_part.element
    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), "42")
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract_num.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "360")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "360")
    indent.set(qn("w:hanging"), "180")
    p_pr.append(indent)
    level.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    r_pr.append(fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2370B5")
    r_pr.append(color)
    level.append(r_pr)
    abstract_num.append(level)
    numbering.append(abstract_num)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), "42")
    abstract_id = OxmlElement("w:abstractNumId")
    abstract_id.set(qn("w:val"), "42")
    num.append(abstract_id)
    numbering.append(num)


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="CV Bullet")
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "42")
    num_pr.append(ilvl)
    num_pr.append(num_id)
    paragraph._p.get_or_add_pPr().insert(0, num_pr)
    run = paragraph.add_run(text)
    set_run_font(run, size=9.3, color=INK)


def add_skill_line(doc, label, value):
    paragraph = doc.add_paragraph(style="CV Compact")
    run = paragraph.add_run(f"{label} : ")
    set_run_font(run, size=9.2, color=NAVY, bold=True)
    run = paragraph.add_run(value)
    set_run_font(run, size=9.2, color=INK)


def build_cv():
    document = Document()
    document.core_properties.title = "CV Lucas Germe - Développeur full-stack junior & AI Builder"
    document.core_properties.subject = "Candidature CDI en développement full-stack"
    document.core_properties.author = "Lucas Germe"

    # Named override cv_a4_compact: French A4 resume, 14 mm margins, one-page target.
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.15)
    section.left_margin = Cm(1.45)
    section.right_margin = Cm(1.45)
    section.header_distance = Cm(0.6)
    section.footer_distance = Cm(0.6)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(9.7)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(3.8)
    normal.paragraph_format.line_spacing = 1.08

    for name, before, after in (
        ("CV Section", 9.0, 3.8),
        ("CV Role", 2.4, 1.7),
        ("CV Bullet", 0, 1.7),
        ("CV Compact", 0, 1.8),
    ):
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = normal
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.03
    styles["CV Bullet"].paragraph_format.left_indent = Cm(0.62)
    styles["CV Bullet"].paragraph_format.first_line_indent = Cm(-0.32)
    create_numbering(document)

    # customer_pack-inspired, ATS-friendly single-column title block.
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(1)
    run = title.add_run("LUCAS GERME")
    set_run_font(run, size=22, color=NAVY, bold=True)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(3)
    run = subtitle.add_run("DÉVELOPPEUR FULL-STACK JUNIOR & AI BUILDER")
    set_run_font(run, size=11.2, color=BLUE, bold=True)

    contact = document.add_paragraph()
    contact.paragraph_format.space_after = Pt(1)
    run = contact.add_run("Thiant - Mobilité France entière")
    set_run_font(run, size=8.9, color=MUTED)
    add_separator(contact)
    add_hyperlink(contact, "06 49 01 05 75", "tel:+33649010575")
    add_separator(contact)
    add_hyperlink(contact, "lucas.germe@gmail.com", "mailto:lucas.germe@gmail.com")

    links = document.add_paragraph()
    links.paragraph_format.space_after = Pt(5)
    add_hyperlink(links, "linkedin.com/in/lucas-germe-5669b9208", "https://www.linkedin.com/in/lucas-germe-5669b9208/")
    add_separator(links)
    add_hyperlink(links, "github.com/germelucas", "https://github.com/germelucas")

    intro = document.add_paragraph()
    intro.paragraph_format.space_after = Pt(4)
    run = intro.add_run(
        "Développeur junior orienté produit, je transforme des besoins métier en applications web utilisables, de l'interface au déploiement. Mon expérience de gérant m'a appris à piloter une activité, prioriser, communiquer avec des clients et résoudre des problèmes concrets. J'utilise les agents IA comme accélérateurs tout en restant responsable du cadrage, des choix, des tests et de la validation."
    )
    set_run_font(run, size=9.5, color=INK)

    add_section_heading(document, "Expérience professionnelle")
    add_role(document, "Gérant", "LG Services / Domicile Clean Cambrai", "Mars 2024 - octobre 2026 (fin prévue)", "Cambrai")
    add_bullet(document, "Création et gestion quotidienne d'une entreprise de services à la personne réalisant environ 70 000 à 83 000 € de chiffre d'affaires annuel.")
    add_bullet(document, "Organisation de 200 à 230 heures de prestations mensuelles, gestion des plannings, des absences et des contraintes clients.")
    add_bullet(document, "Recrutement, intégration et suivi des intervenants ; gestion des contrats clients et salariés.")
    add_bullet(document, "Devis, facturation, administration, relation client, litiges, trésorerie, charges et suivi de la rentabilité.")
    add_bullet(document, "Coordination avec l'expert-comptable, la franchise et les organismes administratifs ; amélioration des processus internes.")

    add_role(document, "Développeur Unity freelance", "ALTAR VISION - ALTAR CORP", "Avril - mai 2026", "6 semaines")
    add_bullet(document, "Seul développeur de la démo Windows jouable d'un FPS rythmique : gameplay, armes, ennemis, score, interface et réglages.")
    add_bullet(document, "Synchronisation des tirs, apparitions, éclairages et retours visuels avec la musique via Unity 6, C#, JSON et Audio DSP.")
    add_bullet(document, "Intégration des assets fournis, débogage, points d'avancement client et livraison du build.")

    add_section_heading(document, "Projets sélectionnés")
    add_role(document, "Planning d'interventions", "Projet full-stack", "Juillet 2026")
    add_bullet(document, "Application React, Node.js/Express et PostgreSQL : vues jour/semaine/mois, glisser-déposer, redimensionnement, prévention des chevauchements et tests automatisés.")
    add_bullet(document, "Architecture API séparant routes, controllers, services et repositories ; déploiement avec Supabase et Vercel.")
    add_role(document, "Applications Android publiées", "Zen Sleep & GOOD", "Octobre 2025 - mars 2026")
    add_bullet(document, "Deux cycles produit complets en Flutter puis React Native/TypeScript : conception, fonctions natives, tests, monétisation et publication Google Play.")

    add_section_heading(document, "Compétences")
    add_skill_line(document, "Frontend", "HTML, CSS, JavaScript, React, TypeScript")
    add_skill_line(document, "Backend & données", "Node.js, Express, API REST, SQL, PostgreSQL, SQLite, Supabase")
    add_skill_line(document, "Tests & livraison", "Tests end-to-end, tests de non-régression, Playwright, Git, GitHub, npm, Vercel")
    add_skill_line(document, "Développement assisté par IA", "Codex, agents de développement, cadrage fonctionnel, diagnostic, contrôle des diffs, tests et validation")
    add_skill_line(document, "Notions", "Kotlin, Java, Spring Boot, Docker, CI/CD et sécurité backend")

    add_section_heading(document, "Langues & mobilité")
    paragraph = document.add_paragraph(style="CV Compact")
    paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(5.8))
    paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(11.5))
    for index, (label, value) in enumerate((
        ("Français", "Courant"),
        ("Anglais", "B1-B2 selon le sujet"),
        ("Mobilité", "Permis B - France entière"),
    )):
        if index:
            paragraph.add_run("\t")
        run = paragraph.add_run(f"{label} : ")
        set_run_font(run, size=8.8, color=NAVY, bold=True)
        run = paragraph.add_run(value)
        set_run_font(run, size=8.8, color=INK)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_cv()
